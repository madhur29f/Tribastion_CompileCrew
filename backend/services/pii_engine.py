from presidio_analyzer import AnalyzerEngine, PatternRecognizer, Pattern
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig
import hashlib
import re
import io


# ---------------------------------------------------------------------------
# Initialize Presidio with custom Indian PII recognizers
# ---------------------------------------------------------------------------

# Custom: Aadhaar Number (4 digits, space, 4 digits, space, 4 digits)
aadhaar_pattern = Pattern(
    name="aadhaar_pattern",
    regex=r"\b\d{4}\s\d{4}\s\d{4}\b",
    score=0.9,
)
aadhaar_recognizer = PatternRecognizer(
    supported_entity="AADHAAR_NUMBER",
    patterns=[aadhaar_pattern],
    supported_language="en",
)

# Custom: PAN Number (5 uppercase letters, 4 digits, 1 uppercase letter)
pan_pattern = Pattern(
    name="pan_pattern",
    regex=r"\b[A-Z]{5}[0-9]{4}[A-Z]\b",
    score=0.9,
)
pan_recognizer = PatternRecognizer(
    supported_entity="PAN_NUMBER",
    patterns=[pan_pattern],
    supported_language="en",
)

# Custom: Date of Birth patterns (DD/MM/YYYY, DD-MM-YYYY, DD.MM.YYYY)
dob_pattern_slash = Pattern(
    name="dob_slash",
    regex=r"\b\d{1,2}[/\-.]\d{1,2}[/\-.]\d{2,4}\b",
    score=0.85,
)
dob_pattern_keyword = Pattern(
    name="dob_keyword",
    regex=r"(?:DOB|D\.O\.B|Date of Birth|Born)\s*[:;]?\s*\d{1,2}[/\-.]\d{1,2}[/\-.]\d{2,4}",
    score=0.95,
)
dob_recognizer = PatternRecognizer(
    supported_entity="DATE_OF_BIRTH",
    patterns=[dob_pattern_slash, dob_pattern_keyword],
    supported_language="en",
)

# Custom: Indian name recognizer for ID card context
# Catches names near keywords like "Name", on Aadhaar/PAN cards
indian_name_pattern = Pattern(
    name="indian_name_context",
    regex=r"(?:Name|नाम)\s*[:;]?\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)",
    score=0.85,
)
indian_name_recognizer = PatternRecognizer(
    supported_entity="PERSON",
    patterns=[indian_name_pattern],
    supported_language="en",
)

# Build analyzer with custom recognizers
analyzer = AnalyzerEngine()
analyzer.registry.add_recognizer(aadhaar_recognizer)
analyzer.registry.add_recognizer(pan_recognizer)
analyzer.registry.add_recognizer(dob_recognizer)
analyzer.registry.add_recognizer(indian_name_recognizer)

anonymizer = AnonymizerEngine()


# ---------------------------------------------------------------------------
# Supported entities
# ---------------------------------------------------------------------------
SUPPORTED_ENTITIES = [
    "PERSON",
    "EMAIL_ADDRESS",
    "PHONE_NUMBER",
    "IP_ADDRESS",
    "AADHAAR_NUMBER",
    "PAN_NUMBER",
    "DATE_OF_BIRTH",
    "LOCATION",
]


# ---------------------------------------------------------------------------
# PII counting (for upload scan)
# ---------------------------------------------------------------------------
def count_pii(text: str) -> int:
    """Count total PII entities found in text."""
    results = analyzer.analyze(text=text, entities=SUPPORTED_ENTITIES, language="en")
    return len(results)


# ---------------------------------------------------------------------------
# PII Risk Scoring Matrix (DPDP Act Compliance)
# ---------------------------------------------------------------------------
PII_SCORE_MATRIX = {
    "AADHAAR_NUMBER": 50,       # Critical: unique national ID
    "PAN_NUMBER": 50,           # Critical: financial ID
    "DATE_OF_BIRTH": 30,        # High: identity verification data
    "PERSON": 15,               # Medium: personal name
    "PHONE_NUMBER": 10,         # Medium: contact info
    "EMAIL_ADDRESS": 5,         # Low-medium: contact info
    "IP_ADDRESS": 5,            # Low-medium: network identifier
    "LOCATION": 5,              # Low-medium: geographic data
}

# Classification tiers based on cumulative score
TIER_THRESHOLDS = [
    (75, "Strictly Confidential"),
    (50, "Confidential"),
    (25, "Internal"),
    (0, "Public"),
]


def calculate_risk_score(text: str, has_faces: bool = False) -> tuple[int, str, dict]:
    """
    Calculate a mathematical risk score (0-100) based on PII density and type.

    Returns (score, classification_tier, breakdown_dict).
    """
    results = analyzer.analyze(text=text, entities=SUPPORTED_ENTITIES, language="en")

    breakdown: dict[str, int] = {}
    raw_score = 0

    for result in results:
        entity_type = result.entity_type
        points = PII_SCORE_MATRIX.get(entity_type, 5)
        raw_score += points
        breakdown[entity_type] = breakdown.get(entity_type, 0) + 1

    # Biometric PII: face detection adds 40 points per face
    if has_faces:
        raw_score += 40
        breakdown["BIOMETRIC_FACE"] = 1

    # Cap at 100
    score = min(raw_score, 100)

    # Determine classification tier
    tier = "Public"
    for threshold, tier_name in TIER_THRESHOLDS:
        if score >= threshold:
            tier = tier_name
            break

    return score, tier, breakdown


# ---------------------------------------------------------------------------
# Anonymization operators by method
# ---------------------------------------------------------------------------
def _get_operators(method: str) -> dict:
    """Return Presidio operator configs for the given sanitization method."""
    redact_op = OperatorConfig("replace", {"new_value": "[REDACTED]"})
    mask_op = OperatorConfig("mask", {"type": "mask", "masking_char": "*", "chars_to_mask": 8, "from_end": False})
    hash_op = OperatorConfig("hash", {"hash_type": "sha256"})

    if method == "smart":
        critical_entities = ["AADHAAR_NUMBER", "PAN_NUMBER", "IP_ADDRESS", "DATE_OF_BIRTH"]
        return {
            entity: (redact_op if entity in critical_entities else mask_op)
            for entity in SUPPORTED_ENTITIES
        }

    if method == "redaction":
        op = redact_op
    elif method == "tokenization":
        op = hash_op
    else:  # masking (default)
        op = mask_op
    
    return {entity: op for entity in SUPPORTED_ENTITIES}


# ---------------------------------------------------------------------------
# Text anonymization
# ---------------------------------------------------------------------------
def anonymize_text(text: str, method: str = "masking") -> str:
    """Run Presidio analysis + anonymization on plain text."""
    results = analyzer.analyze(text=text, entities=SUPPORTED_ENTITIES, language="en")
    operators = _get_operators(method)
    anonymized = anonymizer.anonymize(text=text, analyzer_results=results, operators=operators)
    return anonymized.text


# ---------------------------------------------------------------------------
# SQL anonymization using sqlparse
# Properly parses INSERT statements and only sanitizes string values
# without breaking SQL syntax.
# ---------------------------------------------------------------------------
def anonymize_sql(sql_text: str, method: str = "masking") -> str:
    """Parse SQL with sqlparse, sanitize PII inside string literals of
    INSERT/UPDATE statements while preserving SQL syntax."""
    import sqlparse
    from sqlparse.sql import Statement, Parenthesis, Values
    from sqlparse.tokens import String as TString, Punctuation, DML

    parsed = sqlparse.parse(sql_text)
    output_parts: list[str] = []

    for statement in parsed:
        new_tokens: list[str] = []
        is_insert = False
        in_values = False

        for token in statement.flatten():
            # Detect INSERT keyword
            if token.ttype is DML and token.normalized == "INSERT":
                is_insert = True

            # Detect VALUES keyword
            if is_insert and token.ttype is sqlparse.tokens.Keyword and token.normalized == "VALUES":
                in_values = True

            # Sanitize string literals inside VALUES(...)
            if in_values and token.ttype in (TString.Single, TString.Symbol):
                # Strip outer quotes, sanitize the value, re-quote
                raw_value = token.value.strip("'")
                sanitized_value = anonymize_text(raw_value, method)
                new_tokens.append(f"'{sanitized_value}'")
                continue

            new_tokens.append(token.value)

        output_parts.append("".join(new_tokens))

    return "\n".join(output_parts)


# ---------------------------------------------------------------------------
# SQL PII counting — extract string literals from INSERT VALUES and count PII
# ---------------------------------------------------------------------------
def count_pii_in_sql(sql_text: str) -> int:
    """Count PII in string literals of SQL INSERT statements."""
    import sqlparse
    from sqlparse.tokens import String as TString, DML

    total = 0
    parsed = sqlparse.parse(sql_text)

    for statement in parsed:
        is_insert = False
        in_values = False
        for token in statement.flatten():
            if token.ttype is DML and token.normalized == "INSERT":
                is_insert = True
            if is_insert and token.ttype is sqlparse.tokens.Keyword and token.normalized == "VALUES":
                in_values = True
            if in_values and token.ttype in (TString.Single, TString.Symbol):
                raw_value = token.value.strip("'")
                total += count_pii(raw_value)

    return total


# ---------------------------------------------------------------------------
# PDF text extraction using PyMuPDF with bounding box coordinates
# ---------------------------------------------------------------------------
def extract_pdf_text_with_positions(raw_bytes: bytes) -> list[dict]:
    """Extract text spans from PDF with bounding box coordinates.
    Returns list of {text, bbox, page} dicts."""
    import fitz  # PyMuPDF

    doc = fitz.open(stream=raw_bytes, filetype="pdf")
    spans = []

    for page_num, page in enumerate(doc):
        # Use "dict" output for detailed span-level text + bbox
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        for block in blocks:
            if block.get("type") != 0:  # skip image blocks
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if text:
                        spans.append({
                            "text": text,
                            "bbox": span["bbox"],  # (x0, y0, x1, y1)
                            "page": page_num,
                            "font": span.get("font", ""),
                            "size": span.get("size", 0),
                        })

    doc.close()
    return spans


# ---------------------------------------------------------------------------
# PDF anonymization using PyMuPDF bounding boxes
# ---------------------------------------------------------------------------
def anonymize_pdf(raw_bytes: bytes, method: str = "masking") -> bytes:
    """Read PDF, extract text spans with bounding boxes, find PII,
    and redact the precise bounding box areas. Returns sanitized PDF bytes."""
    import fitz  # PyMuPDF

    doc = fitz.open(stream=raw_bytes, filetype="pdf")

    for page_num, page in enumerate(doc):
        # Get text blocks with bounding boxes
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                # Combine all span text in the line for better PII detection
                line_text = "".join(s.get("text", "") for s in line.get("spans", []))
                if not line_text.strip():
                    continue

                # Analyze PII in the full line text
                results = analyzer.analyze(
                    text=line_text, entities=SUPPORTED_ENTITIES, language="en"
                )
                if not results:
                    continue

                # Map each PII result back to the specific span(s) containing it
                for result in results:
                    pii_text = line_text[result.start:result.end]

                    # Search for the PII text within span bounding boxes
                    char_offset = 0
                    for span in line.get("spans", []):
                        span_text = span.get("text", "")
                        span_start = char_offset
                        span_end = char_offset + len(span_text)

                        # Check if this span overlaps with the PII range
                        if span_end > result.start and span_start < result.end:
                            bbox = fitz.Rect(span["bbox"])

                            sanitized = anonymize_text(pii_text, method)
                            if sanitized == "[REDACTED]":
                                page.add_redact_annot(
                                    bbox,
                                    text="[REDACTED]",
                                    fontsize=span.get("size", 10) * 0.8,
                                    fill=(0, 0, 0),
                                    text_color=(1, 1, 1)
                                )
                            else:
                                page.add_redact_annot(
                                    bbox,
                                    text=sanitized,
                                    fontsize=span.get("size", 10) * 0.8,
                                    fill=(1, 1, 1),  # white bg
                                )

                        char_offset = span_end

        page.apply_redactions()

    output = io.BytesIO()
    doc.save(output)
    doc.close()
    return output.getvalue()


# ---------------------------------------------------------------------------
# DOCX anonymization (sanitize paragraph/table text in memory)
# ---------------------------------------------------------------------------
def anonymize_docx(raw_bytes: bytes, method: str = "masking") -> bytes:
    """Read DOCX, sanitize paragraphs and table cells, return bytes."""
    from docx import Document

    doc = Document(io.BytesIO(raw_bytes))

    # Sanitize paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            sanitized = anonymize_text(paragraph.text, method)
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = sanitized

    # Sanitize table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    sanitized = anonymize_text(cell.text, method)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.text = ""
                        if paragraph.runs:
                            paragraph.runs[0].text = sanitized

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


# ---------------------------------------------------------------------------
# Lazy-loaded EasyOCR reader (heavy model; load once)
# ---------------------------------------------------------------------------
_ocr_reader = None


def _get_ocr_reader():
    global _ocr_reader
    if _ocr_reader is None:
        import easyocr
        _ocr_reader = easyocr.Reader(["en"], gpu=False)
    return _ocr_reader


# ---------------------------------------------------------------------------
# Image PII detection — OCR + Presidio
# ---------------------------------------------------------------------------
def _detect_faces(img_array) -> list:
    """Detect faces in an image using OpenCV Haar cascade.
    Returns list of (x, y, w, h) rectangles."""
    import cv2
    import numpy as np

    gray = cv2.cvtColor(np.array(img_array), cv2.COLOR_RGB2GRAY)

    # Try multiple cascade classifiers for better detection
    cascade_paths = [
        cv2.data.haarcascades + "haarcascade_frontalface_default.xml",
        cv2.data.haarcascades + "haarcascade_frontalface_alt2.xml",
        cv2.data.haarcascades + "haarcascade_profileface.xml",
    ]

    all_faces = []
    for cascade_path in cascade_paths:
        try:
            face_cascade = cv2.CascadeClassifier(cascade_path)
            faces = face_cascade.detectMultiScale(
                gray,
                scaleFactor=1.1,
                minNeighbors=4,
                minSize=(30, 30),
            )
            if len(faces) > 0:
                for (x, y, w, h) in faces:
                    all_faces.append((int(x), int(y), int(w), int(h)))
        except Exception:
            continue

    # De-duplicate overlapping faces
    if not all_faces:
        return []
    return _deduplicate_rects(all_faces)


def _deduplicate_rects(rects: list, overlap_thresh=0.5) -> list:
    """Remove overlapping rectangles, keeping the largest."""
    if not rects:
        return []

    # Sort by area (largest first)
    rects = sorted(rects, key=lambda r: r[2] * r[3], reverse=True)
    keep = []

    for rect in rects:
        x1, y1, w1, h1 = rect
        is_dup = False
        for kept in keep:
            x2, y2, w2, h2 = kept
            # Check overlap
            ix = max(0, min(x1 + w1, x2 + w2) - max(x1, x2))
            iy = max(0, min(y1 + h1, y2 + h2) - max(y1, y2))
            intersection = ix * iy
            area1 = w1 * h1
            if area1 > 0 and intersection / area1 > overlap_thresh:
                is_dup = True
                break
        if not is_dup:
            keep.append(rect)

    return keep


def count_pii_in_image(raw_bytes: bytes) -> int:
    """Run OCR on an image, then count PII entities in the extracted text.
    Also counts detected faces as biometric PII."""
    from PIL import Image
    import numpy as np

    img = Image.open(io.BytesIO(raw_bytes)).convert("RGB")
    reader = _get_ocr_reader()
    ocr_results = reader.readtext(np.array(img))

    # Combine all detected text
    full_text = " ".join(entry[1] for entry in ocr_results)
    text_pii = count_pii(full_text) if full_text.strip() else 0

    # Count faces as biometric PII
    faces = _detect_faces(img)
    return text_pii + len(faces)


# ---------------------------------------------------------------------------
# Image anonymization — pixel-level PII removal
# ---------------------------------------------------------------------------
def anonymize_image(raw_bytes: bytes, method: str = "masking") -> tuple[bytes, str]:
    """
    Read an image, run OCR to find text + bounding boxes, detect PII with
    Presidio, detect faces as biometric PII, then paint solid-colour
    rectangles over ALL PII regions — modifying actual pixels.

    Returns (sanitized_image_bytes, extracted_text_summary).
    """
    from PIL import Image, ImageDraw, ImageFont, ImageFilter
    import numpy as np

    img = Image.open(io.BytesIO(raw_bytes)).convert("RGB")
    draw = ImageDraw.Draw(img)

    # Choose fill colour based on method
    if method == "redaction":
        fill_color = (0, 0, 0)         # black
        text_color = (255, 255, 255)    # white text
        label = "[REDACTED]"
    else:
        fill_color = (40, 40, 40)       # dark grey
        text_color = (180, 180, 180)    # light grey text
        label = "***"

    # -----------------------------------------------------------------
    # Step 1: Face detection — redact faces (biometric PII)
    # -----------------------------------------------------------------
    faces = _detect_faces(img)
    for (fx, fy, fw, fh) in faces:
        padding = 5
        face_rect = (
            max(0, fx - padding),
            max(0, fy - padding),
            min(img.width, fx + fw + padding),
            min(img.height, fy + fh + padding),
        )
        # Paint over face pixels with solid fill
        draw.rectangle(face_rect, fill=fill_color)
        # Label it
        try:
            font_size = max(12, fw // 6)
            try:
                font = ImageFont.truetype("arial.ttf", font_size)
            except (OSError, IOError):
                font = ImageFont.load_default()
            draw.text(
                (face_rect[0] + 4, face_rect[1] + 4),
                "[FACE REDACTED]" if method == "redaction" else "[PHOTO]",
                fill=text_color,
                font=font,
            )
        except Exception:
            pass

    # -----------------------------------------------------------------
    # Step 2: OCR + PII text detection
    # -----------------------------------------------------------------
    reader = _get_ocr_reader()
    ocr_results = reader.readtext(np.array(img))

    if not ocr_results:
        output = io.BytesIO()
        img.save(output, format="PNG")
        return output.getvalue(), ""

    # Build a mapping: for each OCR word, store its text and bounding box
    word_entries = []
    full_text_parts = []
    char_offset = 0

    for bbox, text, confidence in ocr_results:
        start = char_offset
        end = char_offset + len(text)
        xs = [pt[0] for pt in bbox]
        ys = [pt[1] for pt in bbox]
        rect = (int(min(xs)), int(min(ys)), int(max(xs)), int(max(ys)))
        word_entries.append({
            "text": text,
            "start": start,
            "end": end,
            "rect": rect,
        })
        full_text_parts.append(text)
        char_offset = end + 1  # +1 for the space separator

    full_text = " ".join(full_text_parts)

    # Run Presidio detection with lowered threshold for better recall
    results = analyzer.analyze(
        text=full_text,
        entities=SUPPORTED_ENTITIES,
        language="en",
        score_threshold=0.3,   # lower threshold for better recall on ID cards
    )

    if not results:
        output = io.BytesIO()
        img.save(output, format="PNG")
        return output.getvalue(), full_text

    # -----------------------------------------------------------------
    # Step 3: Redact identified PII text regions
    # -----------------------------------------------------------------
    for result in results:
        pii_start = result.start
        pii_end = result.end

        for entry in word_entries:
            if entry["end"] > pii_start and entry["start"] < pii_end:
                rect = entry["rect"]
                padding = 4
                padded_rect = (
                    max(0, rect[0] - padding),
                    max(0, rect[1] - padding),
                    min(img.width, rect[2] + padding),
                    min(img.height, rect[3] + padding),
                )
                draw.rectangle(padded_rect, fill=fill_color)

                try:
                    font_size = max(10, (padded_rect[3] - padded_rect[1]) // 2)
                    try:
                        font = ImageFont.truetype("arial.ttf", font_size)
                    except (OSError, IOError):
                        font = ImageFont.load_default()
                    draw.text(
                        (padded_rect[0] + 2, padded_rect[1] + 2),
                        label,
                        fill=text_color,
                        font=font,
                    )
                except Exception:
                    pass

    output = io.BytesIO()
    img.save(output, format="PNG")
    return output.getvalue(), full_text
